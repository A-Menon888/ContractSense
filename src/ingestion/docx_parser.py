"""
DOCX Parser for ContractSense

Handles Microsoft Word documents using python-docx library.
Extracts text with formatting information and paragraph structure.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from typing import List, Optional
from pathlib import Path
import logging

from . import (
    TextSpan, BoundingBox, Paragraph, Page, ProcessedDocument, 
    DocumentType, logger
)


class DOCXParser:
    """Parser for Microsoft Word documents"""
    
    def parse(self, docx_path: Path) -> ProcessedDocument:
        """
        Parse a DOCX document and return structured data
        
        Args:
            docx_path: Path to the DOCX file
            
        Returns:
            ProcessedDocument with extracted text and structure
        """
        logger.info(f"Parsing DOCX: {docx_path}")
        
        doc = ProcessedDocument(
            document_path=str(docx_path),
            document_type=DocumentType.DOCX
        )
        
        try:
            word_doc = Document(str(docx_path))
            
            # Extract document properties
            doc.metadata.update({
                "author": word_doc.core_properties.author or "",
                "title": word_doc.core_properties.title or "",
                "subject": word_doc.core_properties.subject or "",
                "created": str(word_doc.core_properties.created) if word_doc.core_properties.created else "",
                "modified": str(word_doc.core_properties.modified) if word_doc.core_properties.modified else "",
            })
            
            # Process paragraphs
            pages = self._extract_pages(word_doc)
            doc.pages = pages
            
            doc.metadata["total_pages"] = len(pages)
            doc.metadata["total_paragraphs"] = sum(len(page.paragraphs) for page in pages)
            doc.metadata["total_characters"] = len(doc.full_text)
            doc.metadata["extraction_method"] = "python-docx"
            
        except Exception as e:
            logger.error(f"Error parsing DOCX {docx_path}: {e}")
            doc.processing_errors.append(f"DOCX parsing error: {str(e)}")
        
        return doc
    
    def _extract_pages(self, word_doc: Document) -> List[Page]:
        """
        Extract pages from Word document
        
        Note: Word documents don't have explicit page breaks in the API,
        so we create logical pages based on content or use single page
        """
        pages = []
        current_page = Page(page_num=1, width=8.5 * 72, height=11 * 72)  # Standard letter size in points
        
        char_offset = 0
        paragraph_id = 0
        
        for word_paragraph in word_doc.paragraphs:
            if not word_paragraph.text.strip():
                continue
            
            # Check for page breaks (simplified)
            # In reality, page breaks in DOCX are complex
            if self._has_page_break(word_paragraph):
                if current_page.paragraphs:
                    pages.append(current_page)
                    current_page = Page(
                        page_num=len(pages) + 1,
                        width=8.5 * 72,
                        height=11 * 72
                    )
                    paragraph_id = 0
            
            # Extract paragraph with formatting
            paragraph = self._extract_paragraph(word_paragraph, paragraph_id, char_offset)
            if paragraph:
                current_page.paragraphs.append(paragraph)
                char_offset += len(paragraph.text) + 2  # +2 for paragraph break
                paragraph_id += 1
        
        # Add the final page
        if current_page.paragraphs:
            pages.append(current_page)
        
        # If no pages created, create one empty page
        if not pages:
            pages.append(Page(page_num=1, width=8.5 * 72, height=11 * 72))
        
        return pages
    
    def _extract_paragraph(self, word_paragraph, paragraph_id: int, char_offset: int) -> Optional[Paragraph]:
        """Extract a single paragraph with formatting information"""
        
        if not word_paragraph.text.strip():
            return None
        
        spans = []
        current_offset = char_offset
        
        # Process runs within the paragraph
        for run in word_paragraph.runs:
            if not run.text:
                continue
            
            # Extract formatting
            font_size = None
            font_name = None
            is_bold = run.bold if run.bold is not None else False
            is_italic = run.italic if run.italic is not None else False
            
            if run.font.size:
                font_size = float(run.font.size.pt)
            
            if run.font.name:
                font_name = run.font.name
            
            # Create text span
            span = TextSpan(
                text=run.text,
                start_offset=current_offset,
                end_offset=current_offset + len(run.text),
                font_size=font_size,
                font_name=font_name,
                is_bold=is_bold,
                is_italic=is_italic,
                page_num=1,  # Will be updated when we know the actual page
                confidence=1.0
            )
            
            spans.append(span)
            current_offset += len(run.text)
        
        # If no runs found, create a single span from paragraph text
        if not spans and word_paragraph.text.strip():
            span = TextSpan(
                text=word_paragraph.text,
                start_offset=char_offset,
                end_offset=char_offset + len(word_paragraph.text),
                page_num=1,
                confidence=1.0
            )
            spans.append(span)
        
        if not spans:
            return None
        
        # Create paragraph
        paragraph = Paragraph(
            spans=spans,
            paragraph_id=f"para_{paragraph_id}"
        )
        
        # Analyze paragraph structure
        self._analyze_paragraph_structure(paragraph, word_paragraph)
        
        return paragraph
    
    def _analyze_paragraph_structure(self, paragraph: Paragraph, word_paragraph) -> None:
        """Analyze paragraph to detect headings, numbering, etc."""
        
        text = paragraph.text.strip()
        
        # Check Word's built-in styles for heading detection
        style_name = word_paragraph.style.name.lower()
        if 'heading' in style_name:
            paragraph.is_heading = True
            # Extract heading level from style name
            import re
            match = re.search(r'heading\s*(\d+)', style_name)
            if match:
                paragraph.heading_level = int(match.group(1))
            else:
                paragraph.heading_level = 1
        
        # Check for numbering in Word
        if word_paragraph._element.pPr is not None:
            numPr = word_paragraph._element.pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
            if numPr is not None:
                paragraph.is_numbered = True
                # Extract the visible numbering text (simplified)
                numbering_match = re.match(r'^(\d+\.|\([a-z]\)|\([0-9]+\)|[A-Z]\.|\d+\.\d+)', text)
                if numbering_match:
                    paragraph.number_text = numbering_match.group(0)
        
        # Fallback heuristics if Word styles don't help
        if not paragraph.is_heading:
            # Check for common heading patterns
            is_heading = False
            heading_level = 0
            
            # Short text in all caps might be heading
            if text.isupper() and len(text) < 80:
                is_heading = True
                heading_level = 2
            
            # Bold text that's relatively short
            elif any(span.is_bold for span in paragraph.spans) and len(text) < 100:
                is_heading = True
                heading_level = 3
            
            # Numbered sections
            elif self._detect_numbering_patterns(text):
                paragraph.is_numbered = True
                if len(text) < 150:  # Short numbered items are often headings
                    is_heading = True
                    heading_level = 3
            
            paragraph.is_heading = is_heading
            if heading_level > 0:
                paragraph.heading_level = heading_level
    
    def _detect_numbering_patterns(self, text: str) -> bool:
        """Detect various numbering patterns in text"""
        import re
        
        numbering_patterns = [
            r'^\d+\.\s*',  # 1. 2. 3.
            r'^\d+\.\d+\s*',  # 1.1 1.2
            r'^\([a-z]\)\s*',  # (a) (b) (c)
            r'^\([0-9]+\)\s*',  # (1) (2) (3)
            r'^[A-Z]\.\s*',  # A. B. C.
            r'^[IVX]+\.\s*',  # I. II. III. (Roman numerals)
        ]
        
        for pattern in numbering_patterns:
            if re.match(pattern, text):
                return True
        
        return False
    
    def _has_page_break(self, word_paragraph) -> bool:
        """
        Detect if paragraph has a page break
        
        Note: This is a simplified implementation.
        Real page break detection in DOCX is complex.
        """
        # Check for explicit page break
        try:
            for run in word_paragraph.runs:
                if run._element.xml and 'w:br' in run._element.xml and 'w:type="page"' in run._element.xml:
                    return True
        except:
            pass
        
        return False